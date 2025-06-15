import os
import time
import sqlite3
import pickle
import logging
import hashlib
import fitz  # PyMuPDF
import numpy as np
import pandas as pd
from tqdm import tqdm
from PIL import Image
from openai import AzureOpenAI
from typing import List, Dict, Optional
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AgriculturalML:
    def __init__(self, azure_api_key: str, azure_endpoint: str):
        self.client = AzureOpenAI(
            api_key=azure_api_key,
            api_version="2025-01-01-preview",
            azure_endpoint=azure_endpoint,
        )

        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.db_path = "agricultural_ml.db"
        self.image_dir = "extracted_images"
        self.word_path = "themes_and_subthemes.docx"
        os.makedirs(self.image_dir, exist_ok=True)
        self.init_db()
        
        # Cache for processed PDFs
        self.processed_pdfs = set()
        self.load_processed_pdfs()

    def init_db(self):
        """Initialize database with proper schema"""
        with sqlite3.connect(self.db_path) as conn:
            # Create the documents table with ALL required columns
            conn.execute('''
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    pdf_hash TEXT,
                    page_number INTEGER,
                    content TEXT,
                    embedding BLOB,
                    image_paths TEXT,
                    content_length INTEGER,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Create the processed_pdfs table
            conn.execute('''
                CREATE TABLE IF NOT EXISTS processed_pdfs (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    pdf_hash TEXT UNIQUE,
                    filename TEXT,
                    processed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            # Check if pdf_hash column exists in documents table (for migration)
            cursor = conn.execute("PRAGMA table_info(documents)")
            columns = [column[1] for column in cursor.fetchall()]
            
            if 'pdf_hash' not in columns:
                logger.info("Adding pdf_hash column to existing documents table...")
                conn.execute('ALTER TABLE documents ADD COLUMN pdf_hash TEXT')
            
            # Create indexes for faster queries
            conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_pdf_hash ON documents(pdf_hash)
            ''')
            
            conn.execute('''
                CREATE INDEX IF NOT EXISTS idx_content_length ON documents(content_length)
            ''')

    def load_processed_pdfs(self):
        """Load list of already processed PDFs"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute('SELECT pdf_hash FROM processed_pdfs')
            self.processed_pdfs = {row[0] for row in cursor.fetchall()}

    def get_pdf_hash(self, pdf_path: str) -> str:
        """Generate hash for PDF to check if already processed"""
        hasher = hashlib.sha256()
        with open(pdf_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    def is_relevant_content(self, text: str) -> bool:
        """Filter out irrelevant content like headers, footers, page numbers"""
        text = text.strip()
        
        # Skip if too short
        if len(text) < 50:
            return False
            
        # Skip if mostly numbers (likely page numbers, dates, etc.)
        if len(re.findall(r'\d', text)) / len(text) > 0.5:
            return False
            
        # Skip common headers/footers
        skip_patterns = [
            r'^\d+$',  # Just page numbers
            r'^page \d+',  # Page headers
            r'^\d+\s*$',  # Numbers only
            r'^(chapter|section)\s+\d+$',  # Chapter/section headers only
            r'^(figure|table|chart)\s+\d+',  # Figure/table captions only
            r'^\w{1,3}\s*$',  # Very short abbreviations
        ]
        
        for pattern in skip_patterns:
            if re.match(pattern, text.lower()):
                return False
                
        return True

    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split long text into overlapping chunks for better embedding"""
        if len(text) <= chunk_size:
            return [text]
            
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending within the last 200 characters
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size - 200:
                    end = sentence_end + 1
                    
            chunk = text[start:end].strip()
            if chunk and self.is_relevant_content(chunk):
                chunks.append(chunk)
                
            start = end - overlap
            
        return chunks

    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        return self.embedding_model.encode(texts)

    def beautify_text(self, text: str) -> str:
        prompt = f"""
Enhance the following agricultural document content for clarity and professionalism while maintaining all key information:

{text}

Enhanced version:
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=1000
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Beautification error: {e}")
            return text

    def enhance_query(self, query: str) -> str:
        prompt = f"""
You are a helpful assistant improving search queries for agricultural document retrieval. 
Expand the following query with relevant agricultural terms and context:

Original query: {query}

Enhanced query with agricultural context:
"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=200
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Query enhancement error: {e}")
            return query

    def extract_themes_and_write_word(self, text_data: List[Dict], pdf_hash: str):
        """Extract themes only once per PDF"""
        # Check if themes already extracted for this PDF
        if os.path.exists(f"{self.word_path}_{pdf_hash[:8]}.docx"):
            logger.info(f"Themes already extracted for PDF {pdf_hash[:8]}")
            return
            
        prompt_template = """
Extract high-level themes and subthemes based on the following codebook from the agricultural text:

Codebook:
{
    "Farmer Welfare": ["Farmer Incomes", "Nutrition Security", "Health & Life Insurance", "Livelihood Security", "Costs of Cultivation"],
    "Production Support": ["Water Management", "Soil Health Management", "Fertilisers", "Pest Management", "Access to Credit"],
    "Price Support": ["Price Discovery", "Price Stability", "Fair and Assured Price", "Marketing", "Collectivisation"],
    "Risk Management": ["Crop Insurance", "Contract Farming"],
    "Environmental Sustainability": ["Climate Change", "Sustainable Agriculture", "Crop Diversification", "Central Government Schemes", "State Government Schemes"],
    "Financing": ["Union Budget Allocation", "State Budget Allocation", "Private Financing", "Climate Financing"]
}

Text: {text}

Return only relevant themes found in the text in this format:
Theme | Subtheme | Key Information
------|----------|----------------
"""
        
        # Sample text to avoid token limits
        combined_text = "\n\n".join([entry['content'][:500] for entry in text_data[:20]])
        prompt = prompt_template.replace("{text}", combined_text)

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
            lines = content.strip().splitlines()
            rows = [line.split('|') for line in lines if '|' in line and not line.strip().startswith('Theme')]
            clean_rows = [[cell.strip() for cell in row] for row in rows if len(row) >= 3]

            if clean_rows:
                doc = Document()
                doc.add_heading("Extracted Themes and Subthemes", level=1)

                for theme, subtheme, text in clean_rows:
                    doc.add_paragraph(f"\nTheme: {theme}", style='Heading 2')
                    doc.add_paragraph(f"Subtheme: {subtheme}", style='Heading 3')
                    doc.add_paragraph(text)

                doc.save(f"{self.word_path}_{pdf_hash[:8]}.docx")
                logger.info(f"Themes extracted and saved to {self.word_path}_{pdf_hash[:8]}.docx")
            
        except Exception as e:
            logger.error(f"Word generation error: {e}")

    def store_document(self, pdf_hash: str, page_number: int, content: str, image_paths: List[str]):
        """Store document chunks with PDF hash"""
        chunks = self.chunk_text(content)
        
        if not chunks:
            return
            
        embeddings = self.generate_embeddings(chunks)
        image_paths_str = ';'.join(image_paths)
        
        with sqlite3.connect(self.db_path) as conn:
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                embedding_blob = pickle.dumps(embedding)
                conn.execute('''
                    INSERT INTO documents (pdf_hash, page_number, content, embedding, image_paths, content_length)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (pdf_hash, page_number, chunk, embedding_blob, image_paths_str, len(chunk)))

    def process_pdf(self, pdf_path: str, force_reprocess: bool = False):
        """Process PDF only if not already processed"""
        pdf_hash = self.get_pdf_hash(pdf_path)
        
        if not force_reprocess and pdf_hash in self.processed_pdfs:
            logger.info(f"PDF already processed (hash: {pdf_hash[:8]})")
            return pdf_hash
            
        logger.info(f"Processing new PDF (hash: {pdf_hash[:8]})")
        
        doc = fitz.open(pdf_path)
        all_text_data = []

        for page_num in tqdm(range(len(doc)), desc="Processing PDF pages"):
            page = doc[page_num]
            text = page.get_text("text")
            
            # Skip if no meaningful content
            if not self.is_relevant_content(text):
                continue
                
            images = page.get_images(full=True)
            image_paths = []
            
            for img_index, img in enumerate(images):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image_hash = hashlib.sha1(image_bytes).hexdigest()
                    image_filename = os.path.join(
                        self.image_dir, f"{pdf_hash[:8]}_page{page_num+1}_img{img_index+1}_{image_hash[:8]}.{image_ext}"
                    )
                    
                    # Avoid duplicate images
                    if not os.path.exists(image_filename):
                        with open(image_filename, "wb") as f:
                            f.write(image_bytes)
                    image_paths.append(image_filename)
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_index} from page {page_num}: {e}")

            self.store_document(pdf_hash, page_num + 1, text, image_paths)
            all_text_data.append({"page": page_num + 1, "content": text})

        # Mark PDF as processed
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                INSERT OR REPLACE INTO processed_pdfs (pdf_hash, filename)
                VALUES (?, ?)
            ''', (pdf_hash, os.path.basename(pdf_path)))
            
        self.processed_pdfs.add(pdf_hash)

        # Extract themes only once
        if all_text_data:
            self.extract_themes_and_write_word(all_text_data, pdf_hash)
            
        doc.close()
        return pdf_hash

    def similarity_search(self, query: str, top_k: int = 5, similarity_threshold: float = 0.3) -> List[Dict]:
        """Enhanced similarity search with filtering"""
        improved_query = self.enhance_query(query)
        query_embedding = self.generate_embeddings([improved_query])[0]
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute('''
                SELECT id, pdf_hash, page_number, content, embedding, image_paths, content_length 
                FROM documents 
                WHERE content_length > 50
                ORDER BY content_length DESC
            ''')
            
            results = []
            seen_content = set()  # Avoid duplicate content
            
            for doc_id, pdf_hash, page_number, content, embedding_blob, image_paths_str, content_length in cursor.fetchall():
                # Skip if very similar content already seen
                content_hash = hashlib.md5(content.encode()).hexdigest()
                if content_hash in seen_content:
                    continue
                seen_content.add(content_hash)
                
                doc_embedding = pickle.loads(embedding_blob)
                similarity = cosine_similarity([query_embedding], [doc_embedding])[0][0]
                
                # Apply similarity threshold
                if similarity >= similarity_threshold:
                    results.append({
                        'id': doc_id,
                        'pdf_hash': pdf_hash,
                        'page_number': page_number,
                        'content': content,
                        'image_paths': image_paths_str.split(';') if image_paths_str else [],
                        'similarity': float(similarity),
                        'content_length': content_length
                    })
        
        # Sort by similarity and content length
        results.sort(key=lambda x: (x['similarity'], x['content_length']), reverse=True)
        return results[:top_k]

    def search_and_beautify(self, query: str, top_k: int = 5) -> List[Dict]:
        """Search and beautify results with better filtering"""
        results = self.similarity_search(query, top_k)
        beautified_results = []
        
        for result in results:
            beautified_text = self.beautify_text(result['content'])
            beautified_results.append({
                'pdf_hash': result['pdf_hash'],
                'page_number': result['page_number'],
                'original_text': result['content'],
                'beautified_text': beautified_text,
                'image_paths': result['image_paths'],
                'similarity_score': result['similarity']
            })
            
        return beautified_results

    def get_processing_status(self) -> Dict:
        """Get processing statistics"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute('SELECT COUNT(*) FROM processed_pdfs')
            processed_count = cursor.fetchone()[0]
            
            cursor = conn.execute('SELECT COUNT(*) FROM documents')
            document_count = cursor.fetchone()[0]
            
            return {
                'processed_pdfs': processed_count,
                'document_chunks': document_count
            }


# === Streamlit Interface ===
def main():
    st.set_page_config(page_title="Agricultural PDF Processor", page_icon="📄", layout="wide")
    
    st.title("📄 Agricultural PDF Processor")
    st.markdown("Upload PDF files to extract and embed content for intelligent search.")

    # Sidebar for configuration
    with st.sidebar:
        st.header("Configuration")
        azure_api_key = st.text_input("Azure API Key", type="password")
        azure_endpoint = st.text_input("Azure Endpoint")
        
        st.header("Search Settings")
        top_k = st.slider("Number of results", 1, 20, 5)
        similarity_threshold = st.slider("Similarity threshold", 0.0, 1.0, 0.3, 0.1)

    # Main interface
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("PDF Processing")
        uploaded_files = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)
        force_reprocess = st.checkbox("Force reprocess (ignore cache)")
        
        if uploaded_files and azure_api_key and azure_endpoint:
            ml_system = AgriculturalML(azure_api_key, azure_endpoint)
            
            # Show current status
            status = ml_system.get_processing_status()
            st.info(f"Database contains {status['processed_pdfs']} processed PDFs with {status['document_chunks']} document chunks")
            
            if st.button("Process PDFs"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for i, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Processing {uploaded_file.name}...")
                    
                    # Save uploaded file
                    temp_path = f"temp_{uploaded_file.name}"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.read())
                    
                    try:
                        pdf_hash = ml_system.process_pdf(temp_path, force_reprocess)
                        st.success(f"✅ Processed {uploaded_file.name} (hash: {pdf_hash[:8]})")
                    except Exception as e:
                        st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
                    finally:
                        # Clean up temp file
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                    
                    progress_bar.progress((i + 1) / len(uploaded_files))
                
                status_text.text("Processing complete!")
                st.balloons()

    with col2:
        st.header("Search & Query")
        query = st.text_input("Enter your search query", placeholder="e.g., farmer income, crop insurance, water management")
        
        if query and azure_api_key and azure_endpoint:
            ml_system = AgriculturalML(azure_api_key, azure_endpoint)
            
            with st.spinner("Searching..."):
                results = ml_system.search_and_beautify(query, top_k)
            
            if results:
                st.success(f"Found {len(results)} relevant results")
                
                for i, result in enumerate(results, 1):
                    with st.expander(f"📄 Result {i} - Page {result['page_number']} (Similarity: {result['similarity_score']:.3f})"):
                        st.markdown("**Enhanced Content:**")
                        st.markdown(result['beautified_text'])
                        
                        if st.checkbox(f"Show original text {i}", key=f"original_{i}"):
                            st.markdown("**Original Content:**")
                            st.text(result['original_text'])
                        
                        # Display images
                        for image_path in result['image_paths']:
                            if os.path.exists(image_path):
                                st.image(image_path, width=400, caption=f"Image from page {result['page_number']}")
            else:
                st.warning("No relevant results found. Try adjusting your query or lowering the similarity threshold.")

if __name__ == "__main__":
    main()